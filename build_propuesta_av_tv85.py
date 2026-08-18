from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = Path("outputs")
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "propuesta_estabilidad_video_iglesia_tv85.docx"
CURRENT_DIAGRAM = OUT / "diagrama_actual.png"
PROPOSED_DIAGRAM = OUT / "diagrama_propuesto_tv85.png"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F9"
LIGHT_GREEN = "EAF7EA"
LIGHT_GOLD = "FFF4CC"
GRID = "D9E2EC"


def font(size=11, bold=False, color=INK):
    return {"size": Pt(size), "bold": bold, "color": color}


def set_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def fixed_table(table, widths):
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_paragraph(doc, text="", size=11, bold=False, color=INK, italic=False, after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run(r, size=10.7, color=INK)
    return p


def add_linkish_source(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + ": ")
    set_run(r1, size=9.5, bold=True, color=INK)
    r2 = p.add_run(url)
    set_run(r2, size=9.5, color=BLUE)


def draw_box(draw, xy, fill, outline, text, text_fill=(25, 25, 25), w=3):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=w)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 28)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 22)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    total_h = len(lines) * 30
    y = y1 + ((y2 - y1) - total_h) / 2
    for i, line in enumerate(lines):
        f = title_font if i == 0 else body_font
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=f, fill=text_fill)
        y += 30


def draw_arrow(draw, start, end, color, width=8, label=None):
    draw.line([start, end], fill=color, width=width)
    # arrow head
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 24
    for delta in (2.55, -2.55):
        x = x2 - length * math.cos(angle + delta)
        y = y2 - length * math.sin(angle + delta)
        draw.line([(x2, y2), (x, y)], fill=color, width=width)
    if label:
        try:
            f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 18)
        except Exception:
            f = ImageFont.load_default()
        mx = (x1 + x2) / 2
        my = (y1 + y2) / 2 - 28
        draw.rounded_rectangle((mx - 70, my - 16, mx + 70, my + 16), radius=8, fill=(255, 255, 255), outline=(220, 220, 220))
        bbox = draw.textbbox((0, 0), label, font=f)
        draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 10), label, font=f, fill=(65, 65, 65))


def make_diagrams():
    bg = (255, 255, 255)
    img = Image.new("RGB", (1400, 620), bg)
    d = ImageDraw.Draw(img)
    draw_box(d, (50, 250, 260, 370), (217, 239, 255), (46, 116, 181), "PC\nProPresenter")
    draw_box(d, (480, 250, 700, 370), (237, 237, 237), (140, 140, 140), "Splitter\nHDMI")
    draw_box(d, (1070, 70, 1320, 180), (255, 248, 214), (188, 132, 0), "Proyector 1\nmarca/modelo A")
    draw_box(d, (1070, 255, 1320, 365), (255, 248, 214), (188, 132, 0), "Proyector 2\nmarca/modelo B")
    draw_box(d, (1070, 445, 1320, 555), (231, 247, 231), (55, 150, 75), "TV presentador\nStage Display")
    draw_box(d, (470, 445, 700, 555), (231, 247, 231), (55, 150, 75), "Monitor\noperador")
    draw_arrow(d, (260, 310), (480, 310), (70, 130, 200), label="HDMI largo")
    draw_arrow(d, (700, 285), (1070, 125), (210, 110, 60), label="HDMI largo")
    draw_arrow(d, (700, 310), (1070, 310), (210, 110, 60), label="HDMI largo")
    draw_arrow(d, (260, 340), (1070, 500), (70, 160, 90), label="HDMI largo")
    draw_arrow(d, (260, 355), (470, 500), (70, 160, 90), label="HDMI")
    add_title = "Estado actual: conexiones largas por HDMI y splitter"
    try:
        f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 30)
    except Exception:
        f = ImageFont.load_default()
    d.text((50, 35), add_title, font=f, fill=(30, 30, 30))
    CURRENT_DIAGRAM.parent.mkdir(exist_ok=True)
    img.save(CURRENT_DIAGRAM)

    img = Image.new("RGB", (1400, 620), bg)
    d = ImageDraw.Draw(img)
    d.text((50, 35), "Propuesta: señal estable por HDBaseT sobre Cat6A", font=f, fill=(30, 30, 30))
    draw_box(d, (50, 250, 260, 370), (217, 239, 255), (46, 116, 181), "PC\nProPresenter")
    draw_box(d, (430, 230, 760, 390), (232, 242, 252), (46, 116, 181), "Transmisor HDBaseT\n1 entrada / 2 salidas\nEDID fijo 1080p60")
    draw_box(d, (1030, 70, 1320, 180), (231, 247, 231), (55, 150, 75), "Receptor HDBaseT\n+ TV 85 pulg. 1")
    draw_box(d, (1030, 255, 1320, 365), (231, 247, 231), (55, 150, 75), "Receptor HDBaseT\n+ TV 85 pulg. 2")
    draw_box(d, (1030, 445, 1320, 555), (231, 247, 231), (55, 150, 75), "TV presentador\nStage Display")
    draw_box(d, (430, 450, 760, 555), (245, 245, 245), (120, 120, 120), "Monitor operador\nsalida independiente")
    draw_arrow(d, (260, 310), (430, 310), (70, 130, 200), label="HDMI corto")
    draw_arrow(d, (760, 265), (1030, 125), (55, 150, 75), label="Cat6A directo")
    draw_arrow(d, (760, 310), (1030, 310), (55, 150, 75), label="Cat6A directo")
    draw_arrow(d, (260, 345), (1030, 500), (70, 160, 90), label="HDMI directo")
    draw_arrow(d, (260, 355), (430, 500), (105, 105, 105), label="monitor")
    img.save(PROPOSED_DIAGRAM)


def add_cost_table(doc, title, rows, total_note):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=5)
    fixed_table(table, [2.05, 0.55, 1.05, 1.05, 1.8])
    headers = ["Concepto", "Cant.", "Precio unit.", "Subtotal", "Nota"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=9.2, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run(r, size=8.7, color=INK)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_paragraph(doc, total_note, size=10, bold=True, color=DARK_BLUE, after=10)


def build_doc():
    make_diagrams()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    header = sec.header.paragraphs[0]
    header.text = "Propuesta AV | ProPresenter y pantallas 85 pulgadas"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Documento de referencia para cotización y aprobación")
    set_run(fr, size=8.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Propuesta para estabilizar la señal de video")
    set_run(r, size=24, bold=True, color=RGBColor(0, 0, 0))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Uso con ProPresenter en iglesia pequeña | Monitor operador + dos TVs 85 pulgadas + Stage Display")
    set_run(r, size=13, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    fixed_table(meta, [1.4, 5.1])
    for label, value in [
        ("Preparado para", "Equipo de liderazgo / ministerio de alabanza y multimedia"),
        ("Fecha", "Junio de 2026"),
        ("Objetivo", "Reducir fallas de detección, reinicios y pérdida de señal en pantallas principales sin afectar Stage Display"),
    ]:
        row = meta.rows[len([r for r in meta.rows if r.cells[0].text]) if False else 0]
    for i, (label, value) in enumerate([
        ("Preparado para", "Equipo de liderazgo / ministerio de alabanza y multimedia"),
        ("Fecha", "Junio de 2026"),
        ("Objetivo", "Reducir fallas de detección, reinicios y pérdida de señal en pantallas principales sin afectar Stage Display"),
    ]):
        cells = meta.rows[i].cells
        shade_cell(cells[0], LIGHT_BLUE)
        for text, cell, bold in [(label, cells[0], True), (value, cells[1], False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=9.5, bold=bold, color=INK)

    add_heading(doc, "Resumen ejecutivo")
    add_paragraph(
        doc,
        "Actualmente la computadora con ProPresenter alimenta un monitor local, un splitter HDMI para los dos proyectores y una TV del presentador conectada directo a la PC por otro HDMI largo para Stage Display. "
        "El cable HDMI entre la computadora y el splitter también es largo. "
        "El problema principal no parece ser ProPresenter, sino la combinación de varios tramos HDMI largos, un divisor que puede no manejar bien EDID/handshake y proyectores antiguos de marcas o modelos diferentes. "
        "Esto provoca que Windows detecte las pantallas algunas veces sí y otras no, y que la recuperación requiera desconectar equipos o reiniciar la computadora.",
    )
    add_paragraph(
        doc,
        "La recomendación es reemplazar las corridas largas de la señal principal por HDBaseT sobre cable Cat6A directo. "
        "HDBaseT está diseñado para llevar video HDMI a largas distancias de forma más estable. La TV del presentador debe mantenerse como salida independiente de la PC para Stage Display en ProPresenter; no debe conectarse al splitter porque recibiría la misma señal que las pantallas principales.",
    )

    add_heading(doc, "Objetivos de la mejora")
    for item in [
        "Que las dos pantallas principales reciban la misma señal principal de ProPresenter de forma estable.",
        "Reducir reinicios de la PC, desconexiones y tiempo perdido antes del servicio.",
        "Mantener el monitor de operador independiente y la TV del presentador como salida Stage Display directa desde la PC.",
        "Dejar una instalación más ordenada y fácil de diagnosticar.",
        "Evaluar el cambio de los proyectores por dos TVs de 85 pulgadas si el presupuesto lo permite y si el espacio de montaje lo hace viable.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Estado actual")
    add_paragraph(
        doc,
        "La configuración actual depende de HDMI en tramos largos y un splitter. Esto incluye el tramo entre la PC y el splitter, los tramos del splitter hacia los proyectores y el HDMI largo directo desde la PC hacia la TV del presentador. En distancias largas, HDMI puede ser sensible a cable, energía, orden de encendido y negociación EDID entre computadora, splitter y pantallas.",
    )
    doc.add_picture(str(CURRENT_DIAGRAM), width=Inches(6.55))

    add_heading(doc, "Configuración recomendada")
    add_paragraph(
        doc,
        "La propuesta cambia los tramos largos de la señal principal a Cat6A directo entre transmisor y receptores HDBaseT. El HDMI queda solamente en tramos cortos para pantallas principales. La TV del presentador se mantiene separada, conectada directamente a la PC como Stage Display.",
    )
    doc.add_picture(str(PROPOSED_DIAGRAM), width=Inches(6.55))
    for item in [
        "Configurar la salida de ProPresenter/Windows a 1920 x 1080, 60 Hz.",
        "Desactivar HDR y evitar resoluciones mixtas entre pantallas.",
        "Usar Cat6A de cobre sólido; no pasar HDBaseT por switches de red.",
        "Etiquetar cables y fuentes de poder para facilitar diagnóstico.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Propuesta 1: estabilización económica")
    add_paragraph(
        doc,
        "Esta opción conserva los proyectores actuales y reemplaza la parte más problemática de la señal principal: splitter/cableado HDMI largo hacia proyectores. La TV del presentador se conserva como salida directa de Stage Display desde la PC.",
    )
    rows_budget = [
        ("OREI HD12-EX165-K extensor HDMI 1x2 sobre Cat6/7 con EDID", "1", "$3,438.93", "$3,438.93", "Dos salidas espejo para proyectores"),
        ("Cable Cat6A cobre sólido 100 m", "1", "$1,720.80", "$1,720.80", "Para dos corridas largas"),
        ("HDMI cortos certificados + conectores/placas", "1", "$900.00", "$900.00", "Materiales"),
        ("Canaleta, etiquetas, cinchos y montaje", "1", "$600.00", "$600.00", "Materiales, voluntarios instalan"),
    ]
    add_cost_table(
        doc,
        "Costos estimados - opción económica",
        rows_budget,
        "Subtotal estimado de materiales: $6,659.73 MXN. Rango recomendado para aprobación: $7,000 a $8,500 MXN. La instalación se contempla con voluntarios.",
    )

    doc.add_page_break()
    add_heading(doc, "Propuesta 2: opción óptima sin irse a un sistema caro")
    add_paragraph(
        doc,
        "Esta opción mejora la estabilidad de señal para las dos pantallas principales y reemplaza los dos proyectores por dos TVs de 85 pulgadas. La TV del presentador permanece en una salida separada de Stage Display desde la PC.",
    )
    rows_optimal = [
        ("Divisor HDMI 1x2 OREI UHD-PRO102 con EDID", "1", "$559.83", "$559.83", "Divide solo señal principal"),
        ("AV Access HDBaseT 4K60 / 1080p 230 ft", "2", "$2,747.88", "$5,495.76", "Un kit por TV principal"),
        ("Cable Cat6A cobre sólido 100 m", "1", "$1,720.80", "$1,720.80", "Dos corridas directas"),
        ("HDMI cortos certificados + conectores/placas", "1", "$1,200.00", "$1,200.00", "Materiales"),
        ("Smart TV 85 pulgadas 4K", "2", "$18,000.00", "$36,000.00", "Referencia conservadora por unidad"),
        ("Soportes de pared reforzados para 85 pulgadas", "2", "$1,500.00", "$3,000.00", "Verificar peso/VESA"),
    ]
    add_cost_table(
        doc,
        "Costos estimados - opción óptima",
        rows_optimal,
        "Subtotal con dos TVs de 85 pulgadas: $47,975.39 MXN. Si se conservan los proyectores actuales, el subtotal baja a aproximadamente $8,976.39 MXN. La instalación se contempla con voluntarios.",
    )

    add_heading(doc, "Por qué considerar TVs de 85 pulgadas", level=2)
    add_paragraph(
        doc,
        "La opción de TVs de 85 pulgadas se propone como alternativa a comprar proyectores nuevos. Para una iglesia pequeña, puede ser más práctica si las pantallas se pueden montar a una altura correcta y si el tamaño se ve bien desde las últimas filas. No depende de una marca específica; lo importante es comprar dos TVs iguales, 4K, de marca reconocida, con buena garantía y suficiente brillo para el lugar.",
    )
    for item in [
        "Sin lámparas ni filtros de proyector: reduce mantenimiento periódico y pérdida gradual de brillo por uso de lámpara.",
        "Mejor contraste percibido: las letras, fondos y videos suelen verse con negros más sólidos y color más uniforme que en proyectores antiguos.",
        "Más consistencia visual: dos TVs iguales facilitan que ambos lados se vean con el mismo color, tamaño y resolución.",
        "Mejor tolerancia a luz ambiental: en muchos salones pequeños, una TV grande se ve más clara que un proyector cuando hay luces encendidas.",
        "Limitación importante: una TV de 85 pulgadas puede verse más pequeña que una proyección grande; antes de comprar conviene medir distancia de visualización, altura, ángulo y soporte estructural.",
        "Recomendación de compra: elegir dos unidades iguales de Samsung, LG, TCL o Hisense, 4K, 85 pulgadas, con garantía local y soporte VESA compatible.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Comparación rápida")
    table = doc.add_table(rows=1, cols=4)
    fixed_table(table, [1.35, 1.55, 1.75, 1.85])
    for i, h in enumerate(["Criterio", "Opción económica", "Opción óptima", "Comentario"]):
        shade_cell(table.rows[0].cells[i], LIGHT_GRAY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=9.0, bold=True)
    comparisons = [
        ("Estabilidad de señal", "Alta", "Muy alta", "Ambas eliminan HDMI largo de la señal principal; Stage Display queda directo desde la PC."),
        ("Costo inicial", "Bajo/medio", "Alto", "La óptima sube principalmente por las dos TVs de 85 pulgadas y sus soportes."),
        ("Calidad visual", "Depende de equipos actuales", "Más nítida/consistente", "Dos TVs iguales facilitan color, brillo y resolución."),
        ("Facilidad de diagnóstico", "Mejor que hoy", "Mejor", "Cableado etiquetado y equipo dedicado por proyector."),
        ("Recomendación", "Fase 1 inmediata", "Fase 2 si hay presupuesto", "Se puede hacer por etapas."),
    ]
    for row in comparisons:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run(r, size=8.6)

    add_heading(doc, "Recomendación")
    add_paragraph(
        doc,
        "Aprobar primero la opción económica como Fase 1 para estabilizar la señal principal de los dos proyectores. Mantener la TV del presentador conectada directo a la PC como Stage Display. Si después de estabilizar la señal se decide mejorar también la imagen principal, aprobar la Fase 2 para reemplazar los dos proyectores por dos TVs de 85 pulgadas, siempre validando tamaño visible, altura y soporte estructural.",
    )
    add_paragraph(
        doc,
        "Esta ruta evita gastar de más al inicio, pero corrige el punto técnico más probable de la falla: señal HDMI larga y negociación inconsistente entre dispositivos.",
        bold=True,
        color=DARK_BLUE,
    )

    add_heading(doc, "Notas de implementación")
    for item in [
        "Las corridas Cat6A deben ser directas entre transmisor y receptor; no se conectan a router ni switch.",
        "Antes de instalar definitivamente, probar ambas pantallas principales durante al menos 30 minutos con ProPresenter enviando video y letras.",
        "Guardar una configuración fija en Windows/ProPresenter: 1080p, 60 Hz, salida duplicada solamente para pantallas principales.",
        "Configurar en ProPresenter la TV del presentador como Stage Display independiente, no como salida duplicada.",
        "Antes de comprar TVs, medir la distancia de visualización y confirmar que 85 pulgadas sea suficientemente legible desde la última fila.",
        "Los precios son referencia de Amazon México y pueden cambiar por vendedor, disponibilidad, envío o impuestos.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Fuentes de precios de referencia")
    sources = [
        ("OREI HD12-EX165-K 1x2 HDMI sobre Cat6/7", "https://www.amazon.com.mx/dp/B09ZKL6MXQ"),
        ("AV Access HDBaseT extensor HDMI", "https://www.amazon.com.mx/dp/B019MADOS8"),
        ("OREI UHD-PRO102 splitter HDMI 1x2", "https://www.amazon.com.mx/dp/B0891L7XDH"),
        ("Cable Cat6A cobre sólido 100 m", "https://www.amazon.com.mx/dp/B0F3NYKZS4"),
        ("Búsqueda Amazon MX TVs 85 pulgadas", "https://www.amazon.com.mx/s?k=smart+tv+85+pulgadas+4k"),
        ("Búsqueda Amazon MX soportes TV 85 pulgadas", "https://www.amazon.com.mx/s?k=soporte+tv+85+pulgadas"),
    ]
    for label, url in sources:
        add_linkish_source(doc, label, url)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
